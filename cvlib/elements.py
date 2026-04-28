from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

from .styles import COLORS, FONT_SIZES, SPACING, BULLET_INDENT


def add_centered_text(
    doc: Document,
    text: str,
    size: Pt,
    color: RGBColor = None,
    bold: bool = False,
    space_after: Pt = None,
) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if space_after is not None:
        p.paragraph_format.space_after = space_after
    r = p.add_run(text)
    r.font.size = size
    r.font.bold = bold
    if color:
        r.font.color.rgb = color


def add_section_heading(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = SPACING.SECTION_BEFORE
    p.paragraph_format.space_after = SPACING.SECTION_AFTER
    r = p.add_run(text)
    r.font.size = FONT_SIZES.HEADING
    r.font.bold = True
    r.font.color.rgb = COLORS.BLUE

    pPr = p._element.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "24")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), COLORS.BLUE_HEX)
    pBdr.append(bottom)
    pPr.append(pBdr)


def add_justified_paragraph(
    doc: Document,
    text: str,
    size: Pt = None,
    space_after: Pt = None,
) -> None:
    p = doc.add_paragraph(text)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    if space_after is not None:
        p.paragraph_format.space_after = space_after
    for run in p.runs:
        run.font.size = size or FONT_SIZES.BODY


def add_bullet_list(doc: Document, items: list[str], size: Pt = None) -> None:
    for item in items:
        bp = doc.add_paragraph(item, style="List Bullet")
        bp.paragraph_format.space_after = SPACING.BLOCK_AFTER
        bp.paragraph_format.left_indent = BULLET_INDENT
        for run in bp.runs:
            run.font.size = size or FONT_SIZES.BODY


def add_labeled_line(
    doc: Document,
    label: str,
    value: str,
    label_size: Pt = None,
    value_size: Pt = None,
    value_italic: bool = False,
    space_after: Pt = None,
    space_before: Pt = None,
) -> None:
    p = doc.add_paragraph()
    if space_before is not None:
        p.paragraph_format.space_before = space_before
    if space_after is not None:
        p.paragraph_format.space_after = space_after
    r = p.add_run(label)
    r.font.bold = True
    r.font.size = label_size or FONT_SIZES.BODY
    r = p.add_run(value)
    r.font.size = value_size or FONT_SIZES.BODY
    r.font.italic = value_italic


def add_job_header(
    doc: Document,
    title: str,
    company: str,
    dates: str,
) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = SPACING.ITEM_BEFORE
    p.paragraph_format.space_after = SPACING.ITEM_AFTER
    r = p.add_run(title)
    r.font.bold = True
    r.font.size = FONT_SIZES.SUBTITLE
    r.font.color.rgb = COLORS.BLACK
    p.add_run(" | ").font.size = FONT_SIZES.BODY
    r = p.add_run(company)
    r.font.bold = True
    r.font.size = FONT_SIZES.SUBTITLE

    d = doc.add_paragraph(dates)
    d.paragraph_format.space_after = SPACING.BLOCK_AFTER
    d.runs[0].font.size = FONT_SIZES.BODY
    d.runs[0].font.italic = True
    d.runs[0].font.color.rgb = COLORS.GRAY


def add_tech_stack(doc: Document, tech: str, label: str = "Tech Stack") -> None:
    add_labeled_line(
        doc,
        f"{label}: ",
        tech,
        value_italic=True,
        space_before=SPACING.TECH_BEFORE,
        space_after=SPACING.GROUP_AFTER,
    )


def add_education_entry(
    doc: Document,
    degree: str,
    school: str,
    location: str,
    dates: str,
    gpa: str = "",
    coursework: list[str] = None,
    coursework_label: str = "Relevant Coursework",
) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = SPACING.ITEM_AFTER
    r = p.add_run(degree)
    r.font.bold = True
    r.font.size = FONT_SIZES.SUBTITLE
    p.add_run(" | ")
    r = p.add_run(school)
    r.font.italic = True
    detail = f" | {location} | {dates}"
    if gpa:
        detail += f" | GPA: {gpa}"
    r = p.add_run(detail)
    r.font.size = FONT_SIZES.BODY

    if coursework:
        cw = doc.add_paragraph()
        cw.paragraph_format.space_after = SPACING.BLOCK_AFTER
        cw.paragraph_format.left_indent = BULLET_INDENT
        r = cw.add_run(f"{coursework_label}: ")
        r.font.bold = True
        r.font.size = FONT_SIZES.SMALL
        r.font.color.rgb = COLORS.GRAY
        r = cw.add_run(", ".join(coursework))
        r.font.size = FONT_SIZES.SMALL
        r.font.italic = True
        r.font.color.rgb = COLORS.GRAY


def _add_hyperlink(paragraph, url: str, text: str, color: RGBColor = None, size: Pt = None):
    part = paragraph.part
    r_id = part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)
    new_run = OxmlElement("w:r")
    rPr = OxmlElement("w:rPr")
    c = OxmlElement("w:color")
    c.set(qn("w:val"), str(color) if color else COLORS.BLUE_HEX)
    rPr.append(c)
    u = OxmlElement("w:u")
    u.set(qn("w:val"), "single")
    rPr.append(u)
    if size:
        sz = OxmlElement("w:sz")
        sz.set(qn("w:val"), str(int(size.pt * 2)))
        rPr.append(sz)
    new_run.append(rPr)
    new_run.text = text
    hyperlink.append(new_run)
    paragraph._element.append(hyperlink)


def add_project_entry(
    doc: Document,
    name: str,
    description: str,
    tech: str,
    link: str = "",
    tech_label: str = "Tech",
    link_label: str = "link",
) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = SPACING.ITEM_AFTER
    p.paragraph_format.space_after = SPACING.ITEM_AFTER
    r = p.add_run(name)
    r.font.bold = True
    r.font.size = FONT_SIZES.BODY
    if link:
        r = p.add_run(" [")
        r.font.size = FONT_SIZES.SMALL
        _add_hyperlink(p, link, link_label, COLORS.BLUE, FONT_SIZES.SMALL)
        r = p.add_run("]")
        r.font.size = FONT_SIZES.SMALL
    r = p.add_run(f" -- {description}")
    r.font.size = FONT_SIZES.BODY

    t = doc.add_paragraph(f"{tech_label}: {tech}")
    t.paragraph_format.space_after = SPACING.BLOCK_AFTER
    t.paragraph_format.left_indent = BULLET_INDENT
    t.runs[0].font.size = FONT_SIZES.SMALL
    t.runs[0].font.italic = True
    t.runs[0].font.color.rgb = COLORS.GRAY


def add_titled_entry(
    doc: Document,
    title: str,
    description: str,
    date: str,
    link: str = "",
    link_label: str = "link",
) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = SPACING.ITEM_BEFORE
    p.paragraph_format.space_after = SPACING.ITEM_AFTER
    p.paragraph_format.left_indent = BULLET_INDENT
    
    r = p.add_run(title)
    r.font.bold = True
    r.font.size = FONT_SIZES.BODY

    if link:
        r = p.add_run(" [")
        r.font.size = FONT_SIZES.SMALL
        _add_hyperlink(p, link, link_label, COLORS.BLUE, FONT_SIZES.SMALL)
        r = p.add_run("]")
        r.font.size = FONT_SIZES.SMALL
    
    if date:
        r = p.add_run(f" - {date}")
        r.font.size = FONT_SIZES.BODY
        r.font.italic = True
        r.font.color.rgb = COLORS.GRAY
    
    if description:
        p2 = doc.add_paragraph(description)
        p2.paragraph_format.left_indent = BULLET_INDENT
        p2.paragraph_format.space_after = SPACING.BLOCK_AFTER
        for run in p2.runs:
            run.font.size = FONT_SIZES.BODY


def add_achievement_entry(
    doc: Document,
    title: str,
    description: str,
    date: str,
) -> None:
    add_titled_entry(doc, title, description, date)
